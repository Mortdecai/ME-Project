import sys
import subprocess
import string

from os                     import path, rename
from docxcompose.composer   import Composer
from docx                   import Document
from pathlib                import Path

try:
    from comtypes import client
    import docx2pdf
except ImportError:
    # system is linux
    client = None


root = str(path.dirname(path.realpath(__file__)))


class Specifications:

    def open(self, file_path):
        try:
            file = open(file_path)
            lines = file.read().splitlines()
            file.close()

            for line in lines:
                pair = line.split('=', 1)
                self.__dict[pair[0].split()[0]] = pair[1].split()[0]

            self.source_dir                 = path.abspath(self.__dict['SOURCE_DIR'])
            self.target_dir                 = path.abspath(self.__dict['TARGET_DIR'])
            self.doc_file_path              = path.join(self.target_dir, self.__dict['DOC_FILE_NAME'])
            self.pdf_file_path              = path.join(self.target_dir, self.__dict['PDF_FILE_NAME'])
            self.doc_file_name              = self.__dict['DOC_FILE_NAME']
            self.pdf_file_name              = self.__dict['PDF_FILE_NAME']
            self.doc_file_type              = self.__dict['DOC_FILE_TYPE']

            self.__dict['SOURCE_DIR']       = self.source_dir
            self.__dict['TARGET_DIR']       = self.target_dir
            self.__dict['DOC_FILE_PATH']    = self.doc_file_path
            self.__dict['PDF_FILE_PATH']    = self.pdf_file_path

        except Exception:
            raise

    def __iter__(self):
        return self.__dict

    def __init__(self, file_path):
        self.__dict = {}
        if file_path:
            self.__open(file_path)

    __open = open


def docx_to_pdf_linux(doc, pdf):
    cmd = 'libreoffice --convert-to pdf'.split() + [doc]
    p = subprocess.Popen(cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE)
    p.wait(timeout=15)

    if path.exists(doc):
        name, ext = doc.split('.')
        rename(name + '.pdf', pdf)

    stdout, stderr = p.communicate()
    if stderr:
        raise subprocess.SubprocessError(stderr)

def docx_to_pdf(doc, pdf):
    doc = path.abspath(doc)
    if client is None:
        return docx_to_pdf_linux(doc, pdf)
    docx2pdf.convert(doc, pdf)

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def enum_documents(directory):
    files = []
    for file in Path(directory).rglob('*.docx'):
        files.append(str(file))
    files.sort()
    return files

def merge_documents(files, output_doc, output_pdf):
    merger = Composer(Document())

    for file in files:
        merger.append(Document(file), False)
        paragraph_count = len(merger.doc.paragraphs)
        if paragraph_count > 0:
            delete_paragraph(merger.doc.paragraphs[paragraph_count - 1])

    merger.save(output_doc)
    if not output_pdf is None:
        docx_to_pdf(output_doc, output_pdf)


def main(argv):
    if len(argv) > 0:
        specs = Specifications(path.abspath(argv[0]))
        merge_documents(enum_documents(specs.source_dir), specs.doc_file_path, specs.pdf_file_path)

main(sys.argv[1:])
